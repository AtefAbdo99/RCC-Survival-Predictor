from flask import Flask, request, render_template, redirect, url_for, send_from_directory, jsonify
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
import base64
import os
from docx import Document
import logging
import sys
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score

app = Flask(__name__)
app.secret_key = 'your_secret_key'

# Set up logging to console
handler = logging.StreamHandler(sys.stdout)
handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
app.logger.addHandler(handler)
app.logger.setLevel(logging.INFO)

# Load the model
model_path = "model/model.pkl"
try:
    loaded_data = joblib.load(model_path)
    app.logger.info("Model file loaded successfully.")
    
    if isinstance(loaded_data, tuple):
        model = loaded_data[0]  # Assume the model is the first element of the tuple
        if len(loaded_data) > 1:
            model_performance = loaded_data[1]  # Assume performance data is the second element
        else:
            model_performance = {'predictions': [], 'total_cases': 0}
    elif isinstance(loaded_data, dict):
        model = loaded_data['model']
        model_performance = loaded_data.get('performance', {'predictions': [], 'total_cases': 0})
    else:
        model = loaded_data
        model_performance = {'predictions': [], 'total_cases': 0}
    
    app.logger.info(f"Model feature names in order: {model.feature_names_in_.tolist()}")
except FileNotFoundError:
    app.logger.error(f"Model file not found at {model_path}")
    model = None
    model_performance = {'predictions': [], 'total_cases': 0}
except Exception as e:
    app.logger.error(f"Error loading model: {str(e)}")
    model = None
    model_performance = {'predictions': [], 'total_cases': 0}

if model is None:
    raise RuntimeError("Failed to load model. Please check the file path and ensure the file exists.")

predictions_log_path = 'predictions_log.csv'

def preprocess_input_data(data):
    input_data = pd.DataFrame([data])
    
    app.logger.info(f"Initial column names: {input_data.columns.tolist()}")
    
    # Strict column mapping to match model's expected feature names
    column_mapping = {
        'insurance-status': 'insurance_status',
        'insurance_status': 'insurance_status',
        'laterality': 'laterality ',  # Add space to match model's expectation
        'Laterality': 'laterality ',
        'LATERALITY': 'laterality ',
        'lymph-vascular-invasion': 'lymph_vascular_invasion',
        'lymph_vascular_invasion': 'lymph_vascular_invasion',
        'regional-nodes-examined': 'regional_nodes_examined',
        'regional_nodes_examined': 'regional_nodes_examined',
        'regional-nodes-positive': 'regional_nodes_positive',
        'regional_nodes_positive': 'regional_nodes_positive',
        'charlson-deyo score': 'charlson-deyo score'
    }
    
    # Rename columns based on the mapping
    input_data.rename(columns=column_mapping, inplace=True)
    
    app.logger.info(f"Column names after renaming: {input_data.columns.tolist()}")
    
    # Ensure all expected columns are present with exact naming and order
    expected_columns = model.feature_names_in_.tolist()
    
    for col in expected_columns:
        if col not in input_data.columns:
            app.logger.warning(f"Adding missing column: {col}")
            input_data[col] = None
    
    # Reorder columns to match the order expected by the model
    input_data = input_data[expected_columns]
    
    app.logger.info(f"Final column names: {input_data.columns.tolist()}")
    app.logger.info(f"Data types: {input_data.dtypes}")
    
    return input_data

def encode_categorical(input_data):
    categorical_cols = {
        'facility': {'community cancer program': 0, 'comprehensive community cancer program': 1, 'academic/research program': 2},
        'sex': {'male': 1, 'female': 0},
        'race': {'white': 0, 'black': 1, 'asian': 2, 'other': 3},
        'ethnicity': {'hispanic': 1, 'non-hispanic': 0},
        'insurance_status': {'insured': 0, 'uninsured': 1},
        'charlson-deyo score': {'0': 0, '1': 1, '2': 2, '3': 3},
        'laterality ': {'right': 0, 'left': 1, 'bilateral': 2},
        'grade': {'low grade': 0, 'intermediate grade': 1, 'high grade': 2},
        'lymph_vascular_invasion': {'yes': 1, 'no': 0},
        'surgery': {'yes': 1, 'no': 0},
        'radiotherapy': {'yes': 1, 'no': 0}
    }
    
    for col, mapping in categorical_cols.items():
        if col in input_data.columns:
            app.logger.info(f"Encoding column: {col}")
            input_data[col] = input_data[col].astype(str).str.lower()
            unknown_categories = set(input_data[col].unique()) - set(mapping.keys()) - {'nan', 'none'}
            if unknown_categories:
                raise ValueError(f"Unknown categories in {col}: {', '.join(unknown_categories)}")
            input_data[col] = input_data[col].map(mapping)
            if input_data[col].isnull().any():
                app.logger.warning(f"Null values found in {col} after encoding. These will be handled later.")
    return input_data

def get_os_description(os_months):
    if os_months < 12:
        return "Poor"
    elif 12 <= os_months < 24:
        return "Fair"
    elif 24 <= os_months < 36:
        return "Good"
    else:
        return "Excellent"

@app.route('/')
def welcome():
    return render_template('welcome.html')

@app.route('/predict', methods=['GET', 'POST'])
def predict():
    if request.method == 'POST':
        try:
            data = request.form.to_dict()
            app.logger.info("Received data: %s", data)
            
            input_data = preprocess_input_data(data)
            app.logger.info("Input data after preprocessing: %s", input_data)
            
            # Encode categorical variables
            input_data = encode_categorical(input_data)
            app.logger.info("Input data after encoding: %s", input_data)
            
            # Convert numeric columns
            numeric_cols = ['age', 'income', 'regional_nodes_positive', 'regional_nodes_examined']
            for col in numeric_cols:
                if col in input_data.columns:
                    input_data[col] = pd.to_numeric(input_data[col], errors='coerce')
                    if input_data[col].isnull().any():
                        app.logger.warning(f"Null values in {col} after conversion to numeric")
                        raise ValueError(f"Invalid numeric value for {col}")
            
            app.logger.info("Input data after numeric conversion: %s", input_data)
            
            # Check for any remaining null values
            if input_data.isnull().values.any():
                null_columns = input_data.columns[input_data.isnull().any()].tolist()
                raise ValueError(f"Null values found in columns: {', '.join(null_columns)}")

            # Ensure input features exactly match model features
            if not np.array_equal(input_data.columns, model.feature_names_in_):
                mismatch = np.where(input_data.columns != model.feature_names_in_)[0]
                error_msg = f"Feature mismatch at indices {mismatch}. Expected: {model.feature_names_in_[mismatch]}, Got: {input_data.columns[mismatch]}"
                app.logger.error(error_msg)
                raise ValueError(error_msg)
            
            # Make prediction
            prediction = model.predict(input_data)[0]
            
            # Update model performance
            if 'predictions' not in model_performance:
                model_performance['predictions'] = []
            
            model_performance['predictions'].append(prediction)
            model_performance['total_cases'] = model_performance.get('total_cases', 0) + 1
            
            # Recalculate performance metrics
            if len(model_performance['predictions']) > 1:
                y_true = model_performance['predictions'][:-1]  # All previous predictions
                y_pred = model_performance['predictions'][1:]   # All predictions shifted by 1
                
                model_performance['mse'] = mean_squared_error(y_true, y_pred)
                model_performance['mae'] = mean_absolute_error(y_true, y_pred)
                model_performance['r2'] = r2_score(y_true, y_pred)
            
            # Save updated model and performance metrics
            joblib.dump((model, model_performance), model_path)
            
            os_description = get_os_description(prediction)

            if not os.path.exists(predictions_log_path):
                input_data.to_csv(predictions_log_path, index=False)
            else:
                input_data.to_csv(predictions_log_path, mode='a', header=False, index=False)
            
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 6))
            
            # Overall Survival Prediction
            ax1.bar(['Predicted OS'], [prediction], color='#007bff')
            ax1.set_title("Predicted Overall Survival (OS)")
            ax1.set_ylabel("Months")
            ax1.set_ylim(0, max(60, prediction * 1.2))
            
            # Feature Importance
            feature_importance = model.feature_importances_
            features = input_data.columns
            importance_df = pd.DataFrame({'feature': features, 'importance': feature_importance})
            importance_df = importance_df.sort_values('importance', ascending=False).head(10)
            
            sns.barplot(x='importance', y='feature', data=importance_df, ax=ax2)
            ax2.set_title("Top 10 Feature Importances")
            ax2.set_xlabel("Importance")
            ax2.set_ylabel("Features")
            
            plt.tight_layout()
            img = BytesIO()
            plt.savefig(img, format='png', dpi=300, bbox_inches='tight')
            img.seek(0)
            plot_url = base64.b64encode(img.getvalue()).decode()
            
            return render_template('result.html', 
                                   prediction=round(prediction, 2), 
                                   mse=round(model_performance['mse'], 4) if 'mse' in model_performance else "Insufficient data",
                                   mae=round(model_performance['mae'], 4) if 'mae' in model_performance else "Insufficient data",
                                   r2=round(model_performance['r2'], 4) if 'r2' in model_performance else "Insufficient data",
                                   os_description=os_description, 
                                   plot_url=plot_url,
                                   total_cases=model_performance['total_cases'])
        except Exception as e:
            app.logger.error("Error in prediction: %s", str(e))
            return render_template('predict.html', error=str(e))
    
    return render_template('predict.html')

@app.route('/export')
def export():
    if os.path.exists(predictions_log_path):
        df = pd.read_csv(predictions_log_path)
        document = Document()
        document.add_heading('RCC Survival Prediction Log', 0)

        table = document.add_table(df.shape[0] + 1, df.shape[1])
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]

        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j).text = str(df.values[i, j])

        export_path = os.path.join('exports', 'rcc_prediction_log.docx')
        os.makedirs('exports', exist_ok=True)
        document.save(export_path)
        return redirect(url_for('download_file', filename='rcc_prediction_log.docx'))
    return "No predictions to export."

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory('exports', filename, as_attachment=True)

@app.route('/api/predict', methods=['POST'])
def api_predict():
    try:
        data = request.json
        input_data = preprocess_input_data(data)
        
        input_data = encode_categorical(input_data)
        
        numeric_cols = ['age', 'income', 'regional_nodes_positive', 'regional_nodes_examined']
        for col in numeric_cols:
            if col in input_data.columns:
                input_data[col] = pd.to_numeric(input_data[col], errors='coerce')
        
        if input_data.isnull().values.any():
            return jsonify({"error": "Invalid input: Please ensure all fields are filled correctly."}), 400

        prediction = model.predict(input_data)[0]
        os_description = get_os_description(prediction)

        # Update model performance (same as in the predict route)
        if 'predictions' not in model_performance:
            model_performance['predictions'] = []
        
        model_performance['predictions'].append(prediction)
        model_performance['total_cases'] = model_performance.get('total_cases', 0) + 1
        
        if len(model_performance['predictions']) > 1:
            y_true = model_performance['predictions'][:-1]
            y_pred = model_performance['predictions'][1:]
            
            model_performance['mse'] = mean_squared_error(y_true, y_pred)
            model_performance['mae'] = mean_absolute_error(y_true, y_pred)
            model_performance['r2'] = r2_score(y_true, y_pred)
        
        joblib.dump((model, model_performance), model_path)

        return jsonify({
            "prediction": round(prediction, 2),
            "os_description": os_description,
            "total_cases": model_performance['total_cases']
        })
    except Exception as e:
        app.logger.error("API Error: %s", str(e))
        return jsonify({"error": str(e)}), 400

if __name__ == "__main__":
    app.run(debug=True)